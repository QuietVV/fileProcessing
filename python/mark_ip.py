import logging
import os
import re
import sys
from docx import Document


# 函数功能：对给定的IP地址进行脱敏处理，将IP地址的中间两部分替换为 *
# 参数说明：
#   - ip：需要进行脱敏处理的IP地址字符串，格式形如 x.x.x.x
# 返回值：脱敏后的IP地址字符串，如果输入的IP格式不符合要求则返回原IP
def ip_masking(ip):
    parts = ip.split('.')
    if len(parts) == 4:
        parts[1] = '*'
        parts[2] = '*'
        return '.'.join(parts)
    else:
        return ip


# 函数功能：对docx文档中的IP地址进行脱敏处理，并保存文档
# 参数说明：
#   - docx：通过python-docx库的Document类加载后的docx文档对象，包含文档的段落、表格等内容结构
#   - file_path：docx文档对应的文件路径，用于保存处理后的文档时确定保存位置
def replace_string_docx(docx, file_path):
    flag = False  # 用于标记文档中是否有IP地址被成功脱敏，初始化为0表示未发现可脱敏的IP
    docx_path = os.path.abspath(file_path)  # 获取文档的完整路径，方便后续打印详细的处理信息

    # 遍历文档中的段落
    for paragraph in docx.paragraphs:
        runs = paragraph.runs
        for i, run in enumerate(runs):
            # 使用正则表达式查找文本中的IP地址
            ips = re.findall(r'(?:\d{1,3}\.){3}\d{1,3}', run.text)
            for ip in ips:
                # 对查找到的IP地址进行脱敏处理
                mask_ip = ip_masking(ip)
                new_text = run.text.replace(ip, mask_ip)
                runs[i].text = new_text
                flag = True  # 若有IP地址被成功脱敏，则将标记置为1

    # 遍历文档中的表格
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs = paragraph.runs
                    for i, run in enumerate(runs):
                        ips = re.findall(r'(?:\d{1,3}\.){3}\d{1,3}', run.text)
                        for ip in ips:
                            mask_ip = ip_masking(ip)
                            new_text = run.text.replace(ip, mask_ip)
                            runs[i].text = new_text
                            flag = True

    if flag:
        print(f"已对文件 {docx_path} 完成IP脱敏")
    docx.save(file_path)  # 保存处理后的docx文档到原文件路径


# 函数功能：递归遍历给定文件夹及其子文件夹下的所有文件和文件夹
# 参数说明：
#   - folder_path：需要进行遍历操作的文件夹路径，为字符串类型
def traverse_files_and_folders(folder_path):
    items = os.listdir(folder_path)
    # 配置日志记录器，设置日志级别为INFO，日志格式包含时间和具体消息内容
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
    for item in items:
        file_path = os.path.join(folder_path, item)
        file_path = os.path.normpath(file_path)  # 规范化路径，防止因不同操作系统的路径表示差异出现问题

        if os.path.isfile(file_path):
            try:
                # 判断文件是否为docx格式
                if file_path.endswith('.docx'):
                    docx = Document(file_path)  # 打开docx文件，获取对应的文档对象
                    replace_string_docx(docx, file_path)  # 调用函数对docx文档中的IP地址进行脱敏处理
            except Exception as e:
                # 如果出现异常，记录错误日志并跳过当前文件的处理
                logging.error(f"遍历文件 {file_path} 时出现错误，错误信息：{e}，跳过该文件。")
        elif os.path.isdir(file_path):
            try:
                sub_file_path = os.path.join(folder_path, item)
                if not os.path.exists(sub_file_path) and not os.path.isdir(sub_file_path):
                    os.makedirs(sub_file_path)
                traverse_files_and_folders(sub_file_path)  # 递归调用，继续遍历子文件夹
            except Exception as e:
                # 如果出现异常，记录错误日志并跳过当前文件夹的遍历
                logging.error(f"遍历文件夹 {file_path} 时出现错误，错误信息：{e}，跳过该文件夹。")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("请在命令行中指定要处理的文件夹路径，例如：python your_script.py /your/folder/path")
        sys.exit(1)
    folder_path = sys.argv[1]
    traverse_files_and_folders(folder_path)


